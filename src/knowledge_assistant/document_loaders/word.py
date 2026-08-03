from pathlib import Path

from docx import Document as WordDocument

from knowledge_assistant.document_loaders.base import DocumentLoader
from knowledge_assistant.document_loaders.common import (
    create_document,
    validate_file,
)
from knowledge_assistant.models import Document


class WordDocumentLoader(DocumentLoader):
    """Extract paragraphs and tables from DOCX files."""

    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx"})

    def load(self, path: Path) -> Document:
        resolved_path = validate_file(path)
        word_document = WordDocument(str(resolved_path))

        sections: list[str] = []

        for paragraph in word_document.paragraphs:
            text = paragraph.text.strip()

            if text:
                sections.append(text)

        for table_index, table in enumerate(
            word_document.tables,
            start=1,
        ):
            table_lines = [f"[Table {table_index}]"]

            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell in row.cells
                ]

                if any(values):
                    table_lines.append(
                        " | ".join(values)
                    )

            if len(table_lines) > 1:
                sections.append(
                    "\n".join(table_lines)
                )

        return create_document(
            path=resolved_path,
            content="\n\n".join(sections),
        )