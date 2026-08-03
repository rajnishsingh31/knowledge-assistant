from pathlib import Path

from docx import Document as WordDocument
from openpyxl import Workbook, worksheet
from pypdf import PdfWriter

from knowledge_assistant.document_loaders.excel import (
    ExcelDocumentLoader,
)
from knowledge_assistant.document_loaders.factory import (
    DocumentLoaderFactory,
)
from knowledge_assistant.document_loaders.text import (
    TextDocumentLoader,
)
from knowledge_assistant.document_loaders.word import (
    WordDocumentLoader,
)

def test_text_loader_reads_markdown(
    tmp_path: Path,
) -> None:
    path = tmp_path / "notes.md"
    path.write_text(
        "# Notes\n\nVector search uses embeddings.",
        encoding="utf-8",
    )

    document = TextDocumentLoader().load(path)

    assert "Vector search uses embeddings." in document.content
    assert document.source_path == path.resolve()


def test_word_loader_extracts_paragraphs_and_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.docx"

    word_document = WordDocument()
    word_document.add_paragraph("Quarterly results")

    table = word_document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"

    word_document.save(path)

    document = WordDocumentLoader().load(path)

    assert "Quarterly results" in document.content
    assert "Metric | Value" in document.content
    assert "Revenue | 100" in document.content


def test_excel_loader_extracts_worksheets(
    tmp_path: Path,
) -> None:
    path = tmp_path / "sample.xlsx"

    workbook = Workbook()
    worksheet  = workbook.active
    worksheet.title = "Metrics"
    worksheet.append(["Metric", "Value"])
    worksheet.append(["Revenue", 100])

    workbook.save(path)
    workbook.close()

    document = ExcelDocumentLoader().load(path)

    assert "[Worksheet: Metrics]" in document.content
    assert "Metric | Value" in document.content
    assert "Revenue | 100" in document.content


def test_factory_selects_loader_by_extension(
    tmp_path: Path,
) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("Local knowledge.", encoding="utf-8")

    factory = DocumentLoaderFactory(
        loaders=[
            TextDocumentLoader(),
            WordDocumentLoader(),
            ExcelDocumentLoader(),
        ]
    )

    document = factory.load(path)

    assert document.content == "Local knowledge."