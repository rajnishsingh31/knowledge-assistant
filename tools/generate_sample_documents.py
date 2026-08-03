"""Generate reproducible multi-format sample documents.

The generated documents cover distinct technical topics so they can be
added to the main knowledge corpus without duplicating the existing
Python, Linux, RAG, retrieval, and agentic AI notes.
"""

from pathlib import Path

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUTPUT_DIRECTORY = Path("documents")


def generate_cloud_security_docx(
    output_directory: Path,
) -> Path:
    """Generate a DOCX file about cloud security foundations."""

    output_path = output_directory / "cloud-security.docx"

    document = WordDocument()

    title = document.add_heading(
        "Cloud Security Foundations",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(
        "Cloud security protects applications, data, identities, "
        "networks, and infrastructure hosted in cloud environments."
    )

    document.add_heading("Least Privilege", level=1)
    document.add_paragraph(
        "Least privilege grants users, applications, and services only "
        "the permissions required to perform their assigned tasks."
    )
    document.add_paragraph(
        "Permissions should be reviewed regularly and removed when they "
        "are no longer required."
    )

    document.add_heading("Managed Identities", level=1)
    document.add_paragraph(
        "Managed identities allow cloud applications to authenticate "
        "without storing passwords, certificates, or access keys in "
        "application code."
    )
    document.add_paragraph(
        "The cloud platform manages the identity lifecycle and provides "
        "tokens when the application accesses supported services."
    )

    document.add_heading("Private Endpoints", level=1)
    document.add_paragraph(
        "A private endpoint exposes a cloud service through a private "
        "IP address inside a virtual network."
    )
    document.add_paragraph(
        "Traffic remains on the private network path instead of reaching "
        "the service through its public endpoint."
    )

    document.add_heading("Network Isolation", level=1)
    document.add_paragraph(
        "Network isolation limits which networks, subnets, services, "
        "and users can reach a protected resource."
    )
    document.add_paragraph(
        "Common controls include firewall rules, private endpoints, "
        "service endpoints, network security groups, and disabled public "
        "network access."
    )

    document.add_heading("Secret Rotation", level=1)
    document.add_paragraph(
        "Secret rotation replaces credentials periodically or after a "
        "suspected compromise."
    )
    document.add_paragraph(
        "Short-lived credentials and automated rotation reduce the risk "
        "created by long-lived passwords and access keys."
    )

    document.add_heading("Audit Logging", level=1)
    document.add_paragraph(
        "Audit logs record security-sensitive activities such as "
        "authentication attempts, permission changes, secret access, "
        "resource modifications, and administrative actions."
    )
    document.add_paragraph(
        "Logs should be protected from alteration and retained long "
        "enough to support investigations and compliance requirements."
    )

    document.add_heading("Security Control Summary", level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Control"
    header_cells[1].text = "Purpose"
    header_cells[2].text = "Example"

    controls = [
        (
            "Least Privilege",
            "Reduce unnecessary permissions",
            "Read-only access instead of administrator access",
        ),
        (
            "Managed Identity",
            "Remove credentials from code",
            "Application authenticates to cloud storage",
        ),
        (
            "Private Endpoint",
            "Keep service access on a private network",
            "Database reachable only from a virtual network",
        ),
        (
            "Secret Rotation",
            "Limit credential lifetime",
            "Automatically rotate application credentials",
        ),
        (
            "Audit Logging",
            "Support detection and investigation",
            "Record role and permission changes",
        ),
    ]

    for control, purpose, example in controls:
        cells = table.add_row().cells
        cells[0].text = control
        cells[1].text = purpose
        cells[2].text = example

    document.save(output_path)

    return output_path


def generate_software_testing_xlsx(
    output_directory: Path,
) -> Path:
    """Generate an XLSX workbook about software testing."""

    output_path = output_directory / "software-testing.xlsx"

    workbook = Workbook()

    test_types_sheet = workbook.active
    test_types_sheet.title = "Test Types"

    test_types_sheet.append(
        [
            "Test Type",
            "Purpose",
            "Example",
            "Typical Scope",
        ]
    )

    test_type_rows = [
        [
            "Unit Test",
            "Validate one function, method, or class in isolation",
            "Validate withdrawal amount rules",
            "Single code unit",
        ],
        [
            "Integration Test",
            "Validate that multiple components work together",
            "Application reads and writes to a database",
            "Component boundary",
        ],
        [
            "End-to-End Test",
            "Validate the complete user workflow",
            "Upload a document, search it, and generate an answer",
            "Complete system",
        ],
        [
            "Regression Test",
            "Detect whether an existing capability was broken",
            "Re-run scenarios for previously fixed defects",
            "Existing features",
        ],
        [
            "Performance Test",
            "Measure latency, throughput, and resource consumption",
            "Measure search response time under load",
            "System performance",
        ],
        [
            "Security Test",
            "Find vulnerabilities and authorization failures",
            "Verify that one tenant cannot access another tenant's data",
            "Security boundary",
        ],
        [
            "Smoke Test",
            "Quickly verify that critical functionality is operational",
            "Application starts and health endpoint succeeds",
            "Critical paths",
        ],
        [
            "Contract Test",
            "Verify compatibility between service interfaces",
            "Validate an API response against its schema",
            "Service interface",
        ],
    ]

    for row in test_type_rows:
        test_types_sheet.append(row)

    quality_metrics_sheet = workbook.create_sheet(
        "Quality Metrics"
    )

    quality_metrics_sheet.append(
        [
            "Metric",
            "Meaning",
            "Desired Direction",
            "Example Interpretation",
        ]
    )

    quality_metric_rows = [
        [
            "Pass Rate",
            "Percentage of executed tests that pass",
            "Higher",
            "98 percent of tests passed",
        ],
        [
            "Code Coverage",
            "Percentage of code executed by tests",
            "Higher with meaningful assertions",
            "80 percent line coverage",
        ],
        [
            "Defect Escape Rate",
            "Defects discovered after release",
            "Lower",
            "Two production defects per release",
        ],
        [
            "Mean Time to Detect",
            "Average time required to identify a defect",
            "Lower",
            "Failure detected within five minutes",
        ],
        [
            "Mean Time to Repair",
            "Average time required to restore correct behavior",
            "Lower",
            "Service repaired within thirty minutes",
        ],
        [
            "Flaky Test Rate",
            "Tests that fail intermittently without a product defect",
            "Lower",
            "One percent intermittent failures",
        ],
        [
            "Automation Rate",
            "Percentage of repeatable tests that are automated",
            "Higher where automation adds value",
            "Seventy percent of regression scenarios automated",
        ],
    ]

    for row in quality_metric_rows:
        quality_metrics_sheet.append(row)

    test_strategy_sheet = workbook.create_sheet(
        "Test Strategy"
    )

    test_strategy_sheet.append(
        [
            "Stage",
            "Primary Tests",
            "Goal",
        ]
    )

    strategy_rows = [
        [
            "Developer Commit",
            "Unit and static-analysis tests",
            "Detect defects before code review",
        ],
        [
            "Pull Request",
            "Unit, integration, contract, and security checks",
            "Validate the proposed change",
        ],
        [
            "Preproduction",
            "End-to-end, regression, and performance tests",
            "Validate production-like behavior",
        ],
        [
            "Production",
            "Synthetic tests, monitoring, and safe probes",
            "Detect real operational failures",
        ],
    ]

    for row in strategy_rows:
        test_strategy_sheet.append(row)

    for worksheet in workbook.worksheets:
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(
                horizontal="center",
                vertical="center",
            )

        worksheet.freeze_panes = "A2"

        for column_cells in worksheet.columns:
            column_letter = column_cells[0].column_letter

            maximum_length = max(
                len(str(cell.value))
                if cell.value is not None
                else 0
                for cell in column_cells
            )

            worksheet.column_dimensions[
                column_letter
            ].width = min(maximum_length + 3, 65)

    workbook.save(output_path)
    workbook.close()

    return output_path


def generate_incident_response_pdf(
    output_directory: Path,
) -> Path:
    """Generate a text-based PDF about incident response."""

    output_path = output_directory / "incident-response.pdf"

    styles = getSampleStyleSheet()

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title="Incident Response Lifecycle",
        author="Knowledge Assistant Sample Generator",
    )

    story = [
        Paragraph(
            "Incident Response Lifecycle",
            styles["Title"],
        ),
        Spacer(1, 12),
        Paragraph(
            "Incident response is a structured process for detecting, "
            "containing, resolving, and learning from operational or "
            "security incidents.",
            styles["BodyText"],
        ),
        Spacer(1, 14),
        Paragraph(
            "1. Preparation",
            styles["Heading2"],
        ),
        Paragraph(
            "Preparation defines monitoring, alerting, ownership, "
            "escalation paths, communication channels, recovery "
            "procedures, and access to diagnostic tools.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "2. Detection",
            styles["Heading2"],
        ),
        Paragraph(
            "Detection identifies abnormal behavior using logs, "
            "metrics, distributed traces, alerts, health probes, "
            "security signals, and user reports.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "3. Triage",
            styles["Heading2"],
        ),
        Paragraph(
            "Triage determines impact, affected users, severity, "
            "scope, recent changes, and the teams required to respond.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "4. Containment",
            styles["Heading2"],
        ),
        Paragraph(
            "Containment limits the blast radius and prevents further "
            "damage. Examples include disabling a faulty deployment, "
            "isolating an affected resource, blocking malicious traffic, "
            "or applying temporary resource limits.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "5. Eradication",
            styles["Heading2"],
        ),
        Paragraph(
            "Eradication removes the root cause, malicious artifacts, "
            "invalid credentials, faulty configuration, or defective "
            "software responsible for the incident.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "6. Recovery",
            styles["Heading2"],
        ),
        Paragraph(
            "Recovery restores normal service safely. Responders "
            "validate system health, monitor for recurrence, and return "
            "temporarily disabled features in controlled stages.",
            styles["BodyText"],
        ),
        Spacer(1, 10),
        Paragraph(
            "7. Post-Incident Review",
            styles["Heading2"],
        ),
        Paragraph(
            "The post-incident review documents impact, timeline, root "
            "cause, contributing factors, detection gaps, corrective "
            "actions, owners, and target completion dates.",
            styles["BodyText"],
        ),
        Spacer(1, 16),
        Paragraph(
            "Incident Metrics",
            styles["Heading2"],
        ),
    ]

    metric_rows = [
        ["Metric", "Meaning"],
        [
            "Time to Detect",
            "Time between incident start and detection",
        ],
        [
            "Time to Acknowledge",
            "Time between alert generation and responder acknowledgement",
        ],
        [
            "Time to Mitigate",
            "Time required to reduce customer or system impact",
        ],
        [
            "Time to Recover",
            "Time required to restore normal service",
        ],
        [
            "Recurrence Rate",
            "Frequency at which similar incidents happen again",
        ],
        [
            "Customer Impact",
            "Number of users, tenants, or requests affected",
        ],
    ]

    metric_table = Table(
        metric_rows,
        colWidths=[
            1.8 * inch,
            4.7 * inch,
        ],
        repeatRows=1,
    )

    metric_table.setStyle(
        TableStyle(
            [
                (
                    "FONTNAME",
                    (0, 0),
                    (-1, 0),
                    "Helvetica-Bold",
                ),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.5,
                    "black",
                ),
                (
                    "VALIGN",
                    (0, 0),
                    (-1, -1),
                    "TOP",
                ),
                (
                    "LEFTPADDING",
                    (0, 0),
                    (-1, -1),
                    6,
                ),
                (
                    "RIGHTPADDING",
                    (0, 0),
                    (-1, -1),
                    6,
                ),
                (
                    "TOPPADDING",
                    (0, 0),
                    (-1, -1),
                    6,
                ),
                (
                    "BOTTOMPADDING",
                    (0, 0),
                    (-1, -1),
                    6,
                ),
            ]
        )
    )

    story.extend(
        [
            metric_table,
            Spacer(1, 16),
            Paragraph(
                "Effective Incident Response Practices",
                styles["Heading2"],
            ),
        ]
    )

    practices = [
        "Assign one incident commander to coordinate the response.",
        "Maintain a shared timeline of facts, decisions, and actions.",
        "Separate confirmed facts from hypotheses.",
        "Use reversible mitigation steps where possible.",
        "Communicate status at a predictable cadence.",
        "Track corrective actions until completion.",
        "Improve alerts, dashboards, and runbooks after each incident.",
    ]

    story.append(
        ListFlowable(
            [
                ListItem(
                    Paragraph(
                        practice,
                        styles["BodyText"],
                    )
                )
                for practice in practices
            ],
            bulletType="bullet",
        )
    )

    document.build(story)

    return output_path


def main() -> None:
    """Generate all multi-format sample documents."""

    OUTPUT_DIRECTORY.mkdir(
        parents=True,
        exist_ok=True,
    )

    generated_paths = [
        generate_cloud_security_docx(
            OUTPUT_DIRECTORY
        ),
        generate_software_testing_xlsx(
            OUTPUT_DIRECTORY
        ),
        generate_incident_response_pdf(
            OUTPUT_DIRECTORY
        ),
    ]

    print("Sample documents generated:")

    for path in generated_paths:
        print(f"- {path}")


if __name__ == "__main__":
    main()